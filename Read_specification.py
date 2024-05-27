import docx
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from Read_marker_valve import Dereference_Marker
from Cost_calculation import CostFile


class Specification:
    _items = []  # хранит классы товаров и None если он его не знает

    # Спецификация должна быть в формате docx и сама таблица с товарами должна быть первой
    def __init__(self, spec_path=None) -> None:
        self.path = spec_path
        self.table = self.get_table_info()
        self.items_name_quant = self.get_items_name_quant()
        self.get_items()

    @property
    def path(self):
        return self._path

    @path.setter
    def path(self, value):
        try:
            if value is None:
                Tk().withdraw()  # we don't want a full GUI, so keep the root window from appearing
                filename = askopenfilename()  # show an "Open" dialog box and return the path to the selected file
                self._path = filename
            else:
                self._path = value.strip("\"")
        except Exception as fuck_up:
            print(fuck_up)

    def find_tables(self):
        try:
            # Нужно найти нужные таблицы по ключевым словам
            doc = docx.Document(self.path)
            table_indeces = []
            table_index_counter = -1
            table_triggers = ["описание", "клапан", "зип", "ступен", "уплотнения", "шток" "прокладка", "кольцо",
                              "ремень"]
            for table in doc.tables:
                table_index_counter += 1
                table_found = False
                try:
                    for row in table.rows:
                        for cell in row.cells:
                            if any(trigger in cell.text.lower() for trigger in table_triggers):
                                table_indeces.append(table_index_counter)
                                table_found = True
                            if table_found:
                                raise StopIteration
                except StopIteration:
                    continue
            return table_indeces
        except Exception as fuck_up:
            print(fuck_up)

    def get_table_info(self):
        try:
            doc = docx.Document(self.path)
            table_indeces = self.find_tables()
            if not table_indeces:
                raise IndentationError("Таблицы нужной тут нет!")
            print(f"Номера таблиц в спецификации: {table_indeces}")
            list_table = []
            for table_index in table_indeces:
                word_table = doc.tables[table_index]
                for row in word_table.rows:
                    list_table_row = []
                    end = False
                    for cell in row.cells:
                        text = " ".join(cell.text.split("\n"))
                        list_table_row.append(text)
                        if "итого" in text.lower():
                            end = True
                    if end:
                        break
                    list_table.append(list_table_row)
            return list_table
        except Exception as dang_it:
            print(dang_it)

    def get_items_name_quant(self):
        """делает список картежей из наименования и количества этого
        наименования количество может быть только целым числом.
        Дробная часть будет автоматически откидываться"""
        try:
            items = []
            header_index = None
            name_index = None
            quantity_index = None
            for row in self.table:
                for cell in row:
                    if "наименование" in cell.lower() and "агв" in cell.lower():
                        name_index = row.index(cell)
                        header_index = self.table.index(row)
                        print(cell)
                        break
                if name_index is not None:
                    break
            for row in self.table:
                for cell in row:
                    if header_index is not None and "кол" in cell.lower() and row.index(cell) > name_index:
                        quantity_index = row.index(cell)
                        print("кол-во " + cell.lower())
                if quantity_index is not None:
                    break
            if name_index is None or quantity_index is None:
                raise IndexError("В таблице нет столца с названием 'Наименование АГВ' или количества!")
            for row in range(header_index + 1, len(self.table)):
                name = str(self.table[row][name_index])
                quantity = ""
                for i in self.table[row][quantity_index]:
                    if i.isdigit():
                        quantity += i
                    else:
                        break
                quantity = float(quantity) if quantity != "" else 0
                items.append((name, quantity))
            print("Кол-во пунктов в спецификации", len(items))
            return items
        except (IndexError, Exception) as dang_it:
            print(dang_it)

    def get_items(self):
        goods = ["Прокладка", "Уплотнение", "Кольцо", "Шайба", "Фильтр", "Клапан"]
        for specification_agv_name in self.items_name_quant:
            item_name = specification_agv_name[0].lower()
            if item_name is None:
                Specification._items.append(None)
                return
            zip_valve_triggers = ["зип", "ремонт"]
            if "клапан" in specification_agv_name[0].lower() and any(
                    trigger in specification_agv_name[0].lower() for trigger in zip_valve_triggers):
                analize_goods = goods.copy()
                analize_goods.remove("Клапан")
                if all(good.lower() not in specification_agv_name[0].lower() for good in analize_goods):
                    Specification._items.append(ZipValve(specification_agv_name[0]))
            elif "клапан" in specification_agv_name[0].lower() and all(
                    trigger not in specification_agv_name[0].lower() for trigger in zip_valve_triggers):
                analize_goods = goods.copy()
                analize_goods.remove("Клапан")
                if all(good.lower() not in specification_agv_name[0].lower() for good in analize_goods):
                    Specification._items.append(Valve(specification_agv_name[0]))
            else:
                Specification._items.append(None)


class Valve:
    def __init__(self, specification_agv_name):
        self.name = "клапан"
        self.marker = self.get_marker(specification_agv_name)
        self.parameters = self.get_parameters()
        self.get_pressure()
        self.reg_device = Valve.there_reg_device(specification_agv_name)
        self.description = self.get_description()
        self.cost = self.cal_cost(int(self.parameters.diameter_num))

    @staticmethod
    def get_marker(specification_agv_name):
        """Убрать все знаки препинанния и по порядку достали маркер из списка слов"""
        try:
            limited_punctuation = "!\"#$%&'()*+,-/:;<=>?@[\]^_`{|}~"
            words_list = specification_agv_name.translate(str.maketrans('', '', limited_punctuation)).split()
            for word in words_list:
                if "agv" == word.lower():
                    marker = "AGV" + ".".join(words_list[words_list.index(word) + 1:words_list.index(word) + 4])
                    return marker
                if "agv" in word.lower() and len(word) > 3:
                    return word
            else:
                raise KeyError("Нет маркера в описании клапана или зип клапана")
        except KeyError:
            return "Нет Маркера"

    def get_parameters(self):
        for item in Specification._items:
            if item is not None and self.marker == item.marker:
                print("сработал перенос давления")
                return item.parameters
        else:
            return Dereference_Marker(self.marker)

    def get_pressure(self):
        if self.parameters.pressure is None:
            if self.parameters.nvk == "Нагнетательный":
                pressure = input(
                    f"Какое давление нагнетания для {self.parameters.compressor} {self.parameters.stage_num} ступени: ")
                self.parameters.pressure = pressure.strip()
                return self.parameters.pressure
            elif self.parameters.nvk == "Всасывающий":
                pressure = input(
                    f"Какое давление всасывания для {self.parameters.compressor} {self.parameters.stage_num} ступени: ")
                self.parameters.pressure = pressure.strip()
                return self.parameters.pressure
            else:
                pressure = input(
                    f"Давление, обр/комб клапан для {self.parameters.compressor} {self.parameters.stage_num} ступени: ")
                self.parameters.pressure = pressure.strip()
                return self.parameters.pressure
        else:
            return self.parameters.pressure

    @staticmethod
    def there_reg_device(specification_agv_name):
        """Разные раскладки английская и русская дя одних и тех же сочетаний"""
        there_is = ["с ру", "c ру", "c pу", "с py", "c py", " ру ", "с регулирующим устройством"]
        if any(i in specification_agv_name.lower() for i in there_is):
            return "с РУ "
        else:
            return ""

    def cal_cost(self, diameter):
        valve_cost = CostFile().get_cost_valve(self.parameters.diameter_num)
        if valve_cost == "нет такого диаметра":
            valve_cost = CostFile().get_cost_valve(85)
        if self.reg_device:
            if diameter < 150:
                valve_cost += 15000
            elif diameter >= 150 and diameter < 250:
                valve_cost += 25000
            elif diameter >= 250:
                valve_cost += 35000
            return valve_cost
        else:
            return valve_cost

    def get_description(self):
        if self.parameters.nvk == "Нагнетательный":
            description = f"{self.parameters.nvk} клапан {self.reg_device}d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление нагнетания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description
        elif self.parameters.nvk == "Всасывающий":
            description = f"{self.parameters.nvk} клапан {self.reg_device}d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление всасывания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description
        else:
            description = f"{self.parameters.nvk} клапан d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление всасывания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description


class ZipValve(Valve):
    def __init__(self, specification_agv_name):
        super().__init__(specification_agv_name)
        self.description = self.get_description()
        self.cost = CostFile().get_cost_zip(self.parameters.diameter_num)
        self.name = "зип клапана"

    def get_description(self):
        if self.parameters.nvk == "Нагнетательный":
            description = f"ЗИП клапана нагнетательного d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление нагнетания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description
        elif self.parameters.nvk == "Всасывающий":
            description = f"ЗИП клапана всасывающего d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление всасывания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description
        else:
            description = f"ЗИП клапана обратного/комбинированного d{self.parameters.diameter_num} мм " \
                          f"{self.parameters.full_name} {self.parameters.stage_num}-й ступени " \
                          f"компрессора {self.parameters.compressor}, " \
                          f"давление всасывания -  {self.parameters.pressure} кгс/см2, " \
                          f"конечный заказчик - {self.parameters.company.company_name}."
            return description


def main():
    Tk().withdraw()  # we don't want a full GUI, so keep the root window from appearing
    filename = askopenfilename()  # show an "Open" dialog box and return the path to the selected file
    spec = Specification(filename)
    # Результат пока в текстовый файлик запишу
    try:
        with open("C:\\Users\\sk\\Desktop\\АГВ Проекты\\АГВ проектыV2\\АвтоПО\\new_file.txt", "w",
                  encoding="utf-8") as file:
            counter = 0
            for i in spec._items:
                if i is None:
                    counter += 1
                    file.write(f"№ {counter} {spec.items_name_quant[counter - 1]} -  Его не знаю еще\n")
                else:
                    counter += 1
                    file.write(f"№ {counter}, {i.description} ({i.name}, {i.marker})\n")
    except Exception as error:
        print(error)


def test():
    while True:
        print(Valve.there_reg_device(input("Описание: ")))
        input()


if __name__ == "__main__":
    while True:
        test()
        input()
